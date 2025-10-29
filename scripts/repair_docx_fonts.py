#!/usr/bin/env python
"""
Repair a DOCX by re-saving and normalizing fonts for readability.
- Sets default font to 'Times New Roman' for runs that don't have an explicit font
- Preserves Consolas for code blocks and Cambria Math for equations if already set

Usage:
  python scripts/repair_docx_fonts.py --doc backend/DECUONGLUANVAN.docx [--out backend/DECUONGLUANVAN_fixed.docx]
"""
from __future__ import annotations

import argparse
from pathlib import Path
from docx import Document  # type: ignore
from docx.text.run import Run  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from typing import Iterable

SAFE_FONT = "Times New Roman"
CODE_FONT = "Consolas"
MATH_FONT = "Cambria Math"


def _set_run_font_safe(r: Run) -> None:
    # Preserve code and math font if already set
    name = (r.font.name or "").strip() if r.font is not None else ""
    if name in (CODE_FONT, MATH_FONT):
        return
    # Force-safe font for all normal text
    r.font.name = SAFE_FONT
    # Also set low-level rFonts for full coverage (ascii, hAnsi, eastAsia, cs)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), SAFE_FONT)
    rFonts.set(qn('w:hAnsi'), SAFE_FONT)
    rFonts.set(qn('w:eastAsia'), SAFE_FONT)
    rFonts.set(qn('w:cs'), SAFE_FONT)


def _walk_table_cells(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            yield cell
            # nested tables
            for nt in cell.tables:
                for ncell in _walk_table_cells(nt):
                    yield ncell


def normalize_fonts(doc: Document) -> None:
    # Update common styles to safe font
    for style_name in (
        "Normal", "Body Text", "List Paragraph",
        "Heading 1", "Heading 2", "Heading 3",
        "List Bullet", "List Number",
        "Caption", "Table of Figures", "Table of Tables",
        "Title",
    ):
        try:
            style = doc.styles[style_name]
            if hasattr(style, 'font'):
                style.font.name = SAFE_FONT
        except Exception:
            pass

    # Update runs in body paragraphs
    for p in doc.paragraphs:
        for r in p.runs:
            _set_run_font_safe(r)

    # Update runs inside tables
    for t in doc.tables:
        for cell in _walk_table_cells(t):
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_font_safe(r)

    # Update headers and footers across sections
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                for r in p.runs:
                    _set_run_font_safe(r)
        if section.footer:
            for p in section.footer.paragraphs:
                for r in p.runs:
                    _set_run_font_safe(r)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--doc", required=True)
    ap.add_argument("--out", required=False)
    args = ap.parse_args()

    inp = Path(args.doc)
    if not inp.exists():
        raise SystemExit(f"Document not found: {inp}")

    out = Path(args.out) if args.out else inp.with_name(inp.stem + "_fixed" + inp.suffix)

    doc = Document(str(inp))
    normalize_fonts(doc)
    doc.save(str(out))
    print(f"Saved repaired document to {out}")
