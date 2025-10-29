#!/usr/bin/env python
"""
Apply a section-targeted font normalization on a DOCX by paragraph index range.
- Forces Times New Roman for all runs within [start_para, end_para] inclusive
- Preserves Consolas (code) and Cambria Math (equations)
- Also sets rFonts (ascii, hAnsi, eastAsia, cs) for robustness
- Optionally prints before/after font distribution for the range

Usage:
  python scripts/repair_docx_section.py --doc backend/DECUONGLUANVAN.docx --start 333 --end 339 --stats
"""
from __future__ import annotations
import argparse
from collections import Counter
from pathlib import Path
from typing import Tuple
from docx import Document  # type: ignore
from docx.text.run import Run  # type: ignore
from docx.oxml.ns import qn  # type: ignore

SAFE_FONT = "Times New Roman"
CODE_FONT = "Consolas"
MATH_FONT = "Cambria Math"


def _set_run_font_safe(r: Run) -> None:
    name = (r.font.name or "").strip() if r.font is not None else ""
    if name in (CODE_FONT, MATH_FONT):
        return
    r.font.name = SAFE_FONT
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), SAFE_FONT)
    rFonts.set(qn('w:hAnsi'), SAFE_FONT)
    rFonts.set(qn('w:eastAsia'), SAFE_FONT)
    rFonts.set(qn('w:cs'), SAFE_FONT)


def _range_stats(doc: Document, start: int, end: int) -> Tuple[int, Counter]:
    fonts = Counter()
    scanned = 0
    for i in range(max(0, start), min(len(doc.paragraphs)-1, end) + 1):
        p = doc.paragraphs[i]
        for r in p.runs:
            name = (r.font.name or '').strip()
            fonts[name] += 1
            scanned += 1
    return scanned, fonts


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--doc", required=True)
    ap.add_argument("--start", type=int, required=True)
    ap.add_argument("--end", type=int, required=True)
    ap.add_argument("--stats", action="store_true")
    ap.add_argument("--out", required=False)
    args = ap.parse_args()

    doc_path = Path(args.doc)
    doc = Document(str(doc_path))

    if args.stats:
        scanned, before = _range_stats(doc, args.start, args.end)
        print(f"Before: runs={scanned}")
        for k, v in before.most_common(10):
            print(f"  {k or '<empty>'}: {v}")

    # Apply font normalization on range
    for i in range(max(0, args.start), min(len(doc.paragraphs)-1, args.end) + 1):
        p = doc.paragraphs[i]
        for r in p.runs:
            _set_run_font_safe(r)

    out = Path(args.out) if args.out else doc_path
    doc.save(str(out))

    if args.stats:
        doc2 = Document(str(out))
        scanned2, after = _range_stats(doc2, args.start, args.end)
        print(f"After: runs={scanned2}")
        for k, v in after.most_common(10):
            print(f"  {k or '<empty>'}: {v}")

    print(f"Saved section-repaired document to {out}")
