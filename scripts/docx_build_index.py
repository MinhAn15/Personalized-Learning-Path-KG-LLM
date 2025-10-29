#!/usr/bin/env python
"""
Build a chapter/appendix index map for a DOCX document by scanning visible headings.
Outputs JSON with ordered sections and their start/end positions in the body flow (paragraphs/tables order).
"""
from __future__ import annotations
import argparse
import json
import re
from pathlib import Path
from docx import Document  # type: ignore

# Patterns to detect sections
PATTERNS = [
    re.compile(r"^\s*chương\s+(\d+)", re.IGNORECASE),
    re.compile(r"^\s*phụ\s*lục\s*([A-Z])", re.IGNORECASE),
    re.compile(r"^\s*(tài\s*liệu\s*tham\s*khảo|bibliography)\s*$", re.IGNORECASE),
]


def iter_body_elems(document):
    """Yield (kind, index, object) for each body child in order: 'paragraph' or 'table'."""
    body = document._element.body
    idx = 0
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            yield ("paragraph", idx, document.paragraphs[len([e for e in body.iterchildren() if e.tag.endswith('}p') and e is not child])])
        elif child.tag.endswith('}tbl'):
            # Tables are not directly indexable via document.tables by body order; skip mapping to object
            yield ("table", idx, None)
        idx += 1


def build_index(doc: Document):
    # Simpler approach: scan paragraphs in order; record paragraph-order indices for headings
    sections = []
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or '').strip()
        if not text:
            continue
        for pat in PATTERNS:
            m = pat.match(text)
            if m:
                title = text
                sections.append({"title": title, "para_index": i})
                break
    # Compute ranges in paragraph indices
    ranges = []
    for i, s in enumerate(sections):
        start = s["para_index"]
        end = sections[i+1]["para_index"] - 1 if i + 1 < len(sections) else len(doc.paragraphs) - 1
        ranges.append({"title": s["title"], "start_para": start, "end_para": end})
    return ranges


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--doc", required=True)
    args = ap.parse_args()

    doc_path = Path(args.doc)
    doc = Document(str(doc_path))
    ranges = build_index(doc)
    print(json.dumps({"sections": ranges}, ensure_ascii=False, indent=2))
