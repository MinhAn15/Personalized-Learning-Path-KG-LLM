#!/usr/bin/env python
"""
Normalize fonts for paragraphs whose text matches a regex (e.g., list of figures entries).
- Forces Times New Roman for matched paragraph runs (preserves Consolas and Cambria Math)
- Also normalizes common styles: Caption, TOC 1-9

Usage:
  python scripts/repair_docx_match.py --doc backend/DECUONGLUANVAN.docx --pattern "(?:Hình\s*)?1\\.1\\.(?:[1-8])" --stats
"""
from __future__ import annotations
import argparse
import re
from collections import Counter
from pathlib import Path
from docx import Document  # type: ignore
from docx.text.run import Run  # type: ignore
from docx.oxml.ns import qn  # type: ignore

SAFE_FONT = "Times New Roman"
CODE_FONT = "Consolas"
MATH_FONT = "Cambria Math"


def set_run_font_safe(r: Run) -> None:
    name = (r.font.name or "").strip() if r.font is not None else ""
    if name in (CODE_FONT, MATH_FONT):
        return
    r.font.name = SAFE_FONT
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), SAFE_FONT)
    rFonts.set(qn('w:hAnsi'), SAFE_FONT)
    rFonts.set(qn('w:eastAsia'), SAFE_FONT)
    rFonts.set(qn('w:cs'), SAFE_FONT)


def normalize_styles(doc: Document) -> None:
    # Caption and TOC styles often drive figure lists
    for style_name in ["Caption"] + [f"TOC {i}" for i in range(1, 10)]:
        try:
            st = doc.styles[style_name]
            if hasattr(st, 'font'):
                st.font.name = SAFE_FONT
        except Exception:
            pass


def collect_matches(doc: Document, pattern: str):
    rx = re.compile(pattern)
    idxs = []
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or '')
        if rx.search(t):
            idxs.append(i)
    return idxs


def stats_for(doc: Document, idxs):
    cnt = Counter()
    total = 0
    for i in idxs:
        p = doc.paragraphs[i]
        for r in p.runs:
            name = (r.font.name or '').strip()
            cnt[name] += 1
            total += 1
    return total, cnt


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--doc", required=True)
    ap.add_argument("--pattern", required=True)
    ap.add_argument("--stats", action="store_true")
    ap.add_argument("--out", required=False)
    args = ap.parse_args()

    doc_path = Path(args.doc)
    doc = Document(str(doc_path))

    normalize_styles(doc)
    idxs = collect_matches(doc, args.pattern)

    if args.stats:
        before_total, before = stats_for(doc, idxs)
        print(f"Found {len(idxs)} paragraphs matching; runs={before_total}")
        for k, v in before.most_common():
            print(f"  {k or '<empty>'}: {v}")

    # Fix fonts for matched paragraphs
    for i in idxs:
        p = doc.paragraphs[i]
        for r in p.runs:
            set_run_font_safe(r)

    out = Path(args.out) if args.out else doc_path
    doc.save(str(out))

    if args.stats:
        doc2 = Document(str(out))
        after_total, after = stats_for(doc2, idxs)
        print(f"After runs={after_total}")
        for k, v in after.most_common():
            print(f"  {k or '<empty>'}: {v}")

    print(f"Saved match-repaired document to {out}")
