#!/usr/bin/env python
"""
Insert a UTF-8 text snippet into a .docx document after an anchor paragraph.

Example:
  python scripts/update_docx_insert.py \
      --doc DECUONGLUANVAN.docx \
      --text backend/docs/updates/prompt01_ch1_1.1.txt \
      --anchor "Bối cảnh nghiên cứu" \
      --occurrence 1 \
      --inplace

If the anchor is not found, the text is appended at the end. When --inplace is
not provided, the script writes a new file with suffix .updated.docx
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.shared import Pt  # type: ignore


def _normalize(s: str) -> str:
    import re
    return re.sub(r"\s+", " ", s or "").strip().lower()


def _parse_blocks(text: str):
    """Parse the input markdown-like text into structured blocks.
    Yields tuples (kind, payload) where kind in {heading, paragraph, equation, table}.
    - heading: payload = heading text
    - paragraph: payload = paragraph text
    - equation: payload = equation string (inside bracket or LaTeX delimiters)
    - table: payload = list of rows (list[list[str]])
    """
    lines = text.splitlines()
    i = 0
    n = len(lines)

    def is_table_line(s: str) -> bool:
        s = s.strip()
        return s.startswith("|") and s.endswith("|") and ("|" in s[1:-1])

    def is_separator_line(s: str) -> bool:
        s = s.strip()
        return set(s.replace("|", "").replace(":", "").replace("-", "").strip()) == set()

    while i < n:
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        # Heading starting with ###
        if line.lstrip().startswith("### "):
            yield ("heading", line.lstrip()[4:].strip())
            i += 1
            continue
        # Equation block: a line that starts with \[ or [ and ends with \] or ]
        stripped = line.strip()
        if (stripped.startswith("\\[") and stripped.endswith("\\]")) or (stripped.startswith("[") and stripped.endswith("]") and len(stripped) > 2):
            eq = stripped
            # remove delimiters
            if eq.startswith("\\[") and eq.endswith("\\]"):
                eq = eq[2:-2].strip()
            elif eq.startswith("[") and eq.endswith("]"):
                eq = eq[1:-1].strip()
            yield ("equation", eq)
            i += 1
            continue
        # Table block
        if is_table_line(line):
            # collect table lines
            tbl_lines = [line]
            i += 1
            while i < n and is_table_line(lines[i]):
                tbl_lines.append(lines[i])
                i += 1
            # Parse header and rows (skip separator if present in second line)
            rows = []
            for idx, row in enumerate(tbl_lines):
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                # Skip separator (second line often ---)
                if idx == 1 and is_separator_line(row):
                    continue
                rows.append(cells)
            yield ("table", rows)
            continue
        # Paragraph: accumulate until blank line or another block
        para_lines = [line]
        i += 1
        while i < n and lines[i].strip() and not lines[i].lstrip().startswith("### ") and not is_table_line(lines[i]) and not ((lines[i].strip().startswith("\\[") and lines[i].strip().endswith("\\]")) or (lines[i].strip().startswith("[") and lines[i].strip().endswith("]"))):
            para_lines.append(lines[i])
            i += 1
        yield ("paragraph", " ".join(l.strip() for l in para_lines).strip())


def _insert_structured_after_anchor(doc: Document, text: str, anchor: str, occurrence: int = 1) -> bool:
    """Insert structured content parsed from text after the target anchor paragraph.
    Returns True if inserted after anchor, False if anchor not found (content appended at end).
    """
    # Find anchor index
    idx = -1
    count = 0
    for i, p in enumerate(doc.paragraphs):
        if anchor.lower() in p.text.lower():
            count += 1
            if count == occurrence:
                idx = i
                break

    if idx < 0 or idx >= len(doc.paragraphs):
        # Append at end
        for kind, payload in _parse_blocks(text):
            if kind == "heading":
                doc.add_paragraph(payload).style = "Heading 3"
            elif kind == "equation":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(payload)
                r.font.name = "Cambria Math"
                r.font.size = Pt(12)
            elif kind == "table":
                rows = payload
                if not rows:
                    continue
                cols = len(rows[0])
                t = doc.add_table(rows=0, cols=cols)
                t.style = "Table Grid"
                # header
                hdr = t.add_row().cells
                for j in range(cols):
                    hdr[j].text = rows[0][j]
                # data
                for data_row in rows[1:]:
                    cells = t.add_row().cells
                    for j in range(cols):
                        cells[j].text = data_row[j] if j < len(data_row) else ""
            else:
                doc.add_paragraph(payload)
        return False

    # Insert after anchor paragraph by moving created elements next to anchor
    anchor_p = doc.paragraphs[idx]
    after_elem = anchor_p._element
    for kind, payload in _parse_blocks(text):
        if kind == "heading":
            p = doc.add_paragraph(payload)
            try:
                p.style = "Heading 3"
            except Exception:
                pass
            after_elem.addnext(p._element)
            after_elem = p._element
        elif kind == "equation":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(payload)
            r.font.name = "Cambria Math"
            r.font.size = Pt(12)
            after_elem.addnext(p._element)
            after_elem = p._element
        elif kind == "table":
            rows = payload
            if not rows:
                continue
            cols = len(rows[0])
            t = doc.add_table(rows=0, cols=cols)
            try:
                t.style = "Table Grid"
            except Exception:
                pass
            hdr = t.add_row().cells
            for j in range(cols):
                hdr[j].text = rows[0][j]
            for data_row in rows[1:]:
                cells = t.add_row().cells
                for j in range(cols):
                    cells[j].text = data_row[j] if j < len(data_row) else ""
            after_elem.addnext(t._element)
            after_elem = t._element
        else:
            p = doc.add_paragraph(payload)
            after_elem.addnext(p._element)
            after_elem = p._element
    return True


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(description="Insert text into a DOCX after a matching anchor paragraph.")
    parser.add_argument("--doc", required=True, help="Path to .docx file to update.")
    parser.add_argument("--text", required=True, help="Path to UTF-8 text file to insert.")
    parser.add_argument("--anchor", default="Bối cảnh nghiên cứu", help="Anchor phrase to search for (case-insensitive).")
    parser.add_argument("--occurrence", type=int, default=1, help="Which occurrence of the anchor to target (1-based).")
    parser.add_argument("--inplace", action="store_true", help="Overwrite the input .docx instead of writing a new file.")
    parser.add_argument("--dedupe", action="store_true", help="Remove duplicate occurrences of the inserted text, keeping the earliest.")

    args = parser.parse_args(argv)

    doc_path = Path(args.doc)
    text_path = Path(args.text)
    if not doc_path.exists():
        print(f"DOCX file not found: {doc_path}")
        return 1
    if not text_path.exists():
        print(f"Text file not found: {text_path}")
        return 1

    # Read text with robust encoding fallback
    encodings = ["utf-8", "utf-8-sig", "cp1252", "latin-1"]
    last_err = None
    for enc in encodings:
        try:
            text = text_path.read_text(encoding=enc)
            break
        except Exception as e:  # pragma: no cover
            last_err = e
            text = None
            continue
    if text is None:
        raise SystemExit(f"Failed to read text file {text_path}: {last_err}")
    doc = Document(str(doc_path))

    # Skip if content already present (simple containment check)
    joined = "\n".join(p.text for p in doc.paragraphs)
    if text.strip() and text.strip() in joined:
        inserted = True
    else:
        inserted = _insert_structured_after_anchor(doc, text, args.anchor, args.occurrence)

    if args.dedupe and text.strip():
        # Prefer heading line as signature key
        lines = [ln for ln in text.splitlines() if ln.strip()]
        heading_line = next((ln.lstrip()[4:].strip() for ln in lines if ln.lstrip().startswith("### ")), None)
        key_source = heading_line or (lines[0] if lines else "")
        norm_key = _normalize(key_source)[:100]
        matches = []
        for i, p in enumerate(doc.paragraphs):
            if norm_key and norm_key in _normalize(p.text):
                matches.append(i)
        if len(matches) > 1:
            # Keep the first, remove the rest (from the end to keep indices stable)
            for i in reversed(matches[1:]):
                elm = doc.paragraphs[i]._element
                parent = elm.getparent()
                parent.remove(elm)

    if args.inplace:
        out_path = doc_path
    else:
        out_path = doc_path.with_name(doc_path.stem + ".updated" + doc_path.suffix)

    doc.save(str(out_path))
    where = "after anchor" if inserted else "at end"
    print(f"Saved {out_path} ({where}).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
